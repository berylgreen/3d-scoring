# -*- coding: utf-8 -*-
import os
import json
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from docx import Document

from core.config import settings

# 外部数据路径 (从 config 读取)
JSON_DATA_PATH = settings.GRADING_RESULTS_JSON
STUDENT_DIR_BASE = settings.WORKS_ROOT_DIR

# 扩展名
VIDEO_EXTENSIONS = settings.VIDEO_EXTENSIONS
IMAGE_EXTENSIONS = [".jpg", ".jpeg", ".png", ".bmp", ".webp"]

IGNORE_DIRS = {'__pycache__', '.git', '.vscode', 'node_modules', 'Library', 'Logs', 'Temp', 'obj', 'Packages', 'build', 'UserSettings', 'ProjectSettings'}

def _fast_rglob(folder_path: str, extensions: List[str]) -> List[Path]:
    """
    Fast recursive search for files with specific extensions,
    ignoring common massive directories like Unity 'Library' or '.git'.
    """
    results = []
    ext_set = set(ext.lower() for ext in extensions)
    
    for root, dirs, files in os.walk(folder_path):
        # Prune ignored directories
        dirs[:] = [d for d in dirs if d not in IGNORE_DIRS]
        
        root_path = Path(root)
        for f in files:
            if Path(f).suffix.lower() in ext_set:
                results.append(root_path / f)
                
    return results

def load_all_targets() -> List[Dict]:
    """
    统加载所有的评分目标。
    在 individual 模式下，加载所有学生。
    在 group 模式下，加载所有小组。
    """
    if not JSON_DATA_PATH.exists():
        # Fallback: 扫描本地文件夹以生成基础数据，确保即使未运行AI评分，网页端依然可以浏览作品
        raw_targets = collect_targets_from_disk()
        targets = []
        for t in raw_targets:
            folder_name = t.get("folder_name", "")
            folder_path = t.get("folder_path", "")
            class_name = t.get("class_name", "")
            
            if settings.GRADING_MODE == "group":
                targets.append({
                    "target_id": folder_name,
                    "name": folder_name,
                    "folder_path": folder_path,
                    "total_score": 0,
                    "grading_result": {},
                    "individuals": []
                })
            else:
                import re
                student_id = ""
                student_name = folder_name
                project_name = folder_name
                
                # 尝试解析文件夹名称如 "124232023001_张三_保护环境"
                match = re.match(r'^(\d+)[_\-\s]+([^_\-\s]+)[_\-\s]*(.*)$', folder_name)
                if match:
                    student_id = match.group(1)
                    student_name = match.group(2)
                    project_name = match.group(3) if match.group(3) else folder_name
                
                targets.append({
                    "target_id": student_id if student_id else folder_name,
                    "name": student_name,
                    "project_name": project_name,
                    "class_name": class_name,
                    "folder_path": folder_path,
                    "total_score": 0,
                    "grading_result": {},
                    "grading_time": ""
                })
        return targets

    try:
        with open(JSON_DATA_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        targets = []
        for item in data:
            group_info = item.get("group_info", {})
            target_name = group_info.get("group_name") or group_info.get("name") or group_info.get("project_name", "")
            if not target_name:
                continue
            
            folder_path = group_info.get("folder_path", "")
            if not folder_path:
                folder_path = item.get("folder_path", "")
            if not folder_path:
                # 兼容旧版的逻辑，查找文件夹
                real_path = _find_group_folder_path(target_name)
                folder_path = str(real_path) if real_path else ""
            
            targets.append({
                "target_id": target_name,
                "name": target_name,
                "folder_path": folder_path,
                "total_score": group_info.get("total_group_score") or group_info.get("total_score", 0),
                "grading_result": item,
                "individuals": item.get("individuals", [])
            })
        return targets
    except Exception as e:
        print(f"Error loading JSON data: {e}")
        return []

def collect_targets_from_disk() -> List[Dict]:
    """扫描 WORKS_ROOT_DIR 目录，收集所有学生或小组"""
    targets = []
    works_root = str(STUDENT_DIR_BASE)
    
    # 因为 config.py 中已经配置 WORKS_ROOT_DIR 为 DATA_DIR / "作品"
    # 如果该目录不存在，说明根目录下没有“作品”文件夹
    if not os.path.isdir(works_root):
        print(f"⚠️ 提示: 未找到名为 '作品' 的文件夹 ({works_root})，将不扫描任何数据。")
        return targets

    for folder in os.listdir(works_root):
        folder_stripped = folder.strip()
        student_path = os.path.join(works_root, folder_stripped)
        
        if not os.path.isdir(student_path):
            continue
            
        # 忽略系统文件夹和无关文件夹
        if folder_stripped in ("result", "thumbnail_cache", "batch_grader", "core", "student_web", "__pycache__", ".vscode", ".git", "config", "providers", "utils"):
            continue
            
        targets.append({
            "class_name": "作品",
            "folder_path": student_path,
            "folder_name": folder_stripped
        })
        
    return targets

def get_target_by_id(target_id: str) -> Optional[Dict]:
    targets = load_all_targets()
    for t in targets:
        if t["target_id"] == target_id:
            return t
    return None

def _find_group_folder_path(group_name: str) -> Optional[Path]:
    candidate = STUDENT_DIR_BASE / group_name
    if candidate.exists() and candidate.is_dir():
        nested = candidate / group_name
        if nested.exists() and nested.is_dir():
            return nested
        return candidate
    return None

def find_video_file(folder_path: str) -> Optional[str]:
    if not folder_path:
        return None
    if not os.path.exists(folder_path):
        return None
    
    candidates = _fast_rglob(folder_path, VIDEO_EXTENSIONS)
    
    if not candidates:
        return None
        
    def sort_key(p):
        name = p.name.lower()
        score = 0
        if "最终" in name: score += 100
        if "render" in name: score += 50
        if "animation" in name: score += 50
        return (score, p.stat().st_size)
        
    candidates.sort(key=sort_key, reverse=True)
    return str(candidates[0])

def find_docx_file(folder_path: str) -> Optional[str]:
    if not folder_path:
        return None
    if not os.path.exists(folder_path):
        return None
    
    docx_files = _fast_rglob(folder_path, [".docx"])
    if not docx_files:
        return None
        
    docx_files.sort(key=lambda x: x.stat().st_size, reverse=True)
    return str(docx_files[0])

def find_thumbnail(folder_path: str) -> Optional[str]:
    if not folder_path:
        return None
    if not os.path.exists(folder_path):
        return None
    
    images = _fast_rglob(folder_path, IMAGE_EXTENSIONS)
    if images:
        return str(images[0])
    return None

def find_effect_images(folder_path: str) -> List[str]:
    images = []
    if not folder_path or not os.path.exists(folder_path):
        return images
    
    all_images = _fast_rglob(folder_path, IMAGE_EXTENSIONS)
        
    all_images.sort(key=lambda x: x.name)
    used_images = set()

    for i in range(4):
        prefix_underscore = f"{i}_"
        prefix_dash = f"{i}-"
        prefix_only = str(i)
        found = False
        selected_path = None
        
        for img_path in all_images:
            if img_path.name.startswith(prefix_underscore):
                selected_path = str(img_path)
                found = True
                break
        
        if not found:
            for img_path in all_images:
                if img_path.name.startswith(prefix_dash):
                    selected_path = str(img_path)
                    found = True
                    break
        
        if not found:
            for img_path in all_images:
                name = img_path.name
                if name.startswith(prefix_only) and (len(name) == 1 or not name[1].isdigit()):
                    selected_path = str(img_path)
                    found = True
                    break
                    
        if not found:
            if i < len(all_images):
                selected_path = str(all_images[i])
                found = True

        if found and selected_path:
            if selected_path in used_images:
                for img_path in all_images:
                    p = str(img_path)
                    if p not in used_images:
                        selected_path = p
                        break
            images.append(selected_path)
            used_images.add(selected_path)
    
    return images

def find_personal_images(folder_path: str, members: List[Dict]) -> Dict[str, str]:
    images = {}
    if not folder_path or not os.path.exists(folder_path):
        return images
    
    all_images = _fast_rglob(folder_path, IMAGE_EXTENSIONS)
        
    for member in members:
        name = member.get("student_name", "").strip()
        if not name:
            continue
            
        for img_path in all_images:
            img_name = img_path.name
            if "个人展示" in img_name and name in img_name:
                images[name] = str(img_path)
                break
                
    return images

def find_max_files(folder_path: str) -> List[Dict[str, str]]:
    if not folder_path:
        return []
    if not os.path.exists(folder_path):
        return []
    
    source_files = _fast_rglob(folder_path, [".max", ".blend", ".ma", ".mb", ".c4d"])
        
    if not source_files:
        return []
    
    source_files.sort(key=lambda x: x.stat().st_size, reverse=True)
    
    result = []
    for f in source_files:
        size_mb = f.stat().st_size / (1024 * 1024)
        result.append({
            "name": f.name,
            "path": str(f),
            "size": f"{size_mb:.1f} MB"
        })
    return result

def extract_docx_content(docx_path: str, as_html: bool = False) -> str:
    """
    提取DOCX文档内容。
    如果是 batch grading 需要纯文本，as_html=False。
    如果是 Web 显示，as_html=True。
    """
    if as_html:
        return _extract_docx_html(docx_path)
    else:
        return _extract_docx_text(docx_path)

def _extract_docx_text(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return ""

def _extract_docx_html(docx_path: str) -> str:
    import base64
    from docx.oxml.ns import qn
    
    try:
        doc = Document(docx_path)
        html_parts = []
        image_cache = {}
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_data = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    mime = content_type if content_type else "image/png"
                    b64_data = base64.b64encode(image_data).decode('utf-8')
                    image_cache[rel.rId] = f"data:{mime};base64,{b64_data}"
        except Exception as e:
            pass
            
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para_html = _process_paragraph(element, image_cache)
                if para_html:
                    html_parts.append(para_html)
            elif element.tag.endswith('tbl'):
                table_html = _process_table(element)
                if table_html:
                    html_parts.append(table_html)
        
        if not html_parts:
            return "<div class='empty-doc'>文档内容为空</div>"
        
        return "\n".join(html_parts)
        
    except Exception as e:
        return f"<div class='error-doc'>无法读取文档: {str(e)}</div>"

def _process_paragraph(para_element, image_cache: dict) -> str:
    from docx.oxml.ns import qn
    html_content = []
    
    for child in para_element.iter():
        if child.tag.endswith('}t'):
            text = child.text or ""
            if text.strip():
                html_content.append(_escape_html(text))
        elif child.tag.endswith('}blip'):
            embed_attr = child.get(qn('r:embed'))
            if embed_attr and embed_attr in image_cache:
                img_src = image_cache[embed_attr]
                html_content.append(f'<img src="{img_src}" class="doc-img max-w-full h-auto my-4 rounded shadow-sm" />')
                
    for drawing in para_element.iter():
        if drawing.tag.endswith('}drawing'):
            for blip in drawing.iter():
                if blip.tag.endswith('}blip'):
                    embed_attr = blip.get(qn('r:embed'))
                    if embed_attr and embed_attr in image_cache:
                        img_tag = f'<img src="{image_cache[embed_attr]}" class="doc-img max-w-full h-auto my-4 rounded shadow-sm" />'
                        if img_tag not in html_content:
                            html_content.append(img_tag)
                            
    content = "".join(html_content).strip()
    if content:
        return f"<p class='my-2 text-gray-700 leading-relaxed'>{content}</p>"
    return ""

def _process_table(table_element) -> str:
    rows_html = []
    for tr in table_element.iter():
        if tr.tag.endswith('}tr'):
            cells_html = []
            for tc in tr.iter():
                if tc.tag.endswith('}tc'):
                    cell_text = []
                    for t in tc.iter():
                        if t.tag.endswith('}t'):
                            if t.text:
                                cell_text.append(_escape_html(t.text))
                    cells_html.append(f"<td class='border border-gray-300 px-4 py-2'>{''.join(cell_text)}</td>")
            if cells_html:
                rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    
    if rows_html:
        return f"<div class='overflow-x-auto my-4'><table class='min-w-full border-collapse border border-gray-300'>{''.join(rows_html)}</table></div>"
    return ""

def _escape_html(text: str) -> str:
    return (text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;"))
